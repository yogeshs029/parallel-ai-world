import io
import re
import logging
from typing import Tuple
import httpx
from bs4 import BeautifulSoup
import pypdf
import docx

logger = logging.getLogger(__name__)

def clean_extracted_text(text: str) -> str:
    """Normalize whitespace, remove excess blank lines and control characters."""
    if not text:
        return ""
    # Replace carriage returns
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Replace multiple spaces with single space
    text = re.sub(r"[ \t]+", " ", text)
    # Replace 3+ consecutive newlines with 2 newlines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract readable text from a PDF file using pypdf."""
    try:
        stream = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(stream)
        pages_text = []
        for i, page in enumerate(reader.pages):
            page_content = page.extract_text() or ""
            if page_content.strip():
                pages_text.append(f"[Page {i + 1}]\n{page_content.strip()}")
        return clean_extracted_text("\n\n".join(pages_text))
    except Exception as e:
        logger.error(f"Failed to extract PDF text: {e}")
        raise ValueError(f"Could not read PDF document: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract paragraphs and table text from a DOCX document."""
    try:
        stream = io.BytesIO(file_bytes)
        doc = docx.Document(stream)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_cells:
                    paragraphs.append(" | ".join(row_cells))
        return clean_extracted_text("\n\n".join(paragraphs))
    except Exception as e:
        logger.error(f"Failed to extract DOCX text: {e}")
        raise ValueError(f"Could not read Word document: {str(e)}")

def extract_text_from_txt_or_md(file_bytes: bytes) -> str:
    """Extract plaintext from UTF-8 / Latin-1 encoded text and markdown."""
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            return clean_extracted_text(file_bytes.decode(enc))
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file with standard encodings.")

async def extract_text_from_url(url: str) -> Tuple[str, str]:
    """
    Fetch a public web page and extract readable article text and title.
    Returns: (clean_text, page_title)
    """
    if not url.startswith("http://") and not url.startswith("https://"):
        url = "https://" + url

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }

    try:
        async with httpx.AsyncClient(timeout=15.0, follow_redirects=True) as client:
            resp = await client.get(url, headers=headers)
            resp.raise_for_status()
            html = resp.text

        soup = BeautifulSoup(html, "html.parser")

        # Remove non-content elements
        for element in soup(["script", "style", "nav", "footer", "header", "aside", "form", "svg", "noscript"]):
            element.decompose()

        # Extract title
        title_tag = soup.find("title")
        page_title = title_tag.get_text().strip() if title_tag else url

        # Extract body text
        body = soup.find("body") or soup
        text = body.get_text(separator="\n")
        clean_text = clean_extracted_text(text)

        if not clean_text or len(clean_text) < 10:
            raise ValueError("Web page contained no readable text content.")

        return clean_text, page_title

    except Exception as e:
        logger.error(f"Failed to fetch or parse URL {url}: {e}")
        raise ValueError(f"We couldn't read this page. Please check the URL or try another.")

def process_file_content(filename: str, file_bytes: bytes) -> Tuple[str, str]:
    """
    Determine file type from filename extension and extract plain text.
    Returns: (extracted_text, mime_type)
    """
    fn_lower = filename.lower()
    if fn_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes), "application/pdf"
    elif fn_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fn_lower.endswith(".md"):
        return extract_text_from_txt_or_md(file_bytes), "text/markdown"
    elif fn_lower.endswith(".txt") or fn_lower.endswith(".json") or fn_lower.endswith(".csv"):
        return extract_text_from_txt_or_md(file_bytes), "text/plain"
    else:
        raise ValueError(f"Unsupported file format. Please upload a PDF, DOCX, TXT, or MD document.")
